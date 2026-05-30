"""Document loading and text extraction for multiple file formats."""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, List

import httpx
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """A document extracted from a file or URL."""
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentLoader:
    """Loads and extracts text from PDFs, DOCX, TXT, and web URLs."""

    @staticmethod
    def _clean_text(text: str) -> str:
        """Normalize whitespace and strip control characters."""
        # Remove control chars except newline/tab
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
        # Collapse runs of whitespace (except newlines) to a single space
        text = re.sub(r"[^\S\n]+", " ", text)
        # Collapse 3+ consecutive newlines to 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def load_pdf(self, file_bytes: bytes, source: str = "pdf") -> List[LoadedDocument]:
        """Extract text from a PDF file.

        Parameters
        ----------
        file_bytes : bytes
            Raw PDF content.
        source : str
            Identifier for the source document.

        Returns
        -------
        list[LoadedDocument]
            One document per page.
        """
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        documents: List[LoadedDocument] = []

        for page_num, page in enumerate(reader.pages, start=1):
            raw = page.extract_text() or ""
            cleaned = self._clean_text(raw)
            if cleaned:
                documents.append(
                    LoadedDocument(
                        text=cleaned,
                        metadata={
                            "source": source,
                            "page_number": page_num,
                            "total_pages": len(reader.pages),
                            "file_type": "pdf",
                        },
                    )
                )

        logger.info("Loaded %d pages from PDF '%s'", len(documents), source)
        return documents

    def load_docx(self, file_bytes: bytes, source: str = "docx") -> List[LoadedDocument]:
        """Extract text from a DOCX file.

        Parameters
        ----------
        file_bytes : bytes
            Raw DOCX content.
        source : str
            Identifier for the source document.

        Returns
        -------
        list[LoadedDocument]
            A single document with the full text content.
        """
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = self._clean_text("\n\n".join(paragraphs))

        if not full_text:
            logger.warning("No text extracted from DOCX '%s'", source)
            return []

        logger.info("Loaded DOCX '%s' (%d characters)", source, len(full_text))
        return [
            LoadedDocument(
                text=full_text,
                metadata={
                    "source": source,
                    "file_type": "docx",
                    "paragraph_count": len(paragraphs),
                },
            )
        ]

    def load_txt(self, file_bytes: bytes, source: str = "txt") -> List[LoadedDocument]:
        """Load a plain-text file.

        Parameters
        ----------
        file_bytes : bytes
            Raw file content.
        source : str
            Identifier for the source document.

        Returns
        -------
        list[LoadedDocument]
        """
        # Try UTF-8 first, fall back to latin-1
        try:
            raw = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            raw = file_bytes.decode("latin-1")

        cleaned = self._clean_text(raw)

        if not cleaned:
            logger.warning("No text in TXT file '%s'", source)
            return []

        logger.info("Loaded TXT '%s' (%d characters)", source, len(cleaned))
        return [
            LoadedDocument(
                text=cleaned,
                metadata={"source": source, "file_type": "txt"},
            )
        ]

    async def load_web_url(self, url: str) -> List[LoadedDocument]:
        """Fetch and extract text from a web page.

        Parameters
        ----------
        url : str
            The URL to fetch.

        Returns
        -------
        list[LoadedDocument]
        """
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            response = await client.get(
                url,
                headers={
                    "User-Agent": "Mozilla/5.0 (Enterprise RAG Assistant Bot)"
                },
            )
            response.raise_for_status()

        soup = BeautifulSoup(response.text, "lxml")

        # Remove script and style elements
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        # Extract title
        title = soup.title.string.strip() if soup.title and soup.title.string else url

        # Get main content
        main = soup.find("main") or soup.find("article") or soup.body or soup
        raw_text = main.get_text(separator="\n", strip=True)
        cleaned = self._clean_text(raw_text)

        if not cleaned:
            logger.warning("No text extracted from URL '%s'", url)
            return []

        logger.info("Loaded URL '%s' (%d characters)", url, len(cleaned))
        return [
            LoadedDocument(
                text=cleaned,
                metadata={
                    "source": url,
                    "title": title,
                    "file_type": "web",
                },
            )
        ]

    def load_file(self, file_bytes: bytes, filename: str) -> List[LoadedDocument]:
        """Auto-detect format and load a file.

        Parameters
        ----------
        file_bytes : bytes
            Raw file content.
        filename : str
            Original file name (used for type detection).

        Returns
        -------
        list[LoadedDocument]

        Raises
        ------
        ValueError
            If the file type is not supported.
        """
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return self.load_pdf(file_bytes, source=filename)
        elif suffix in (".docx",):
            return self.load_docx(file_bytes, source=filename)
        elif suffix in (".txt", ".md", ".csv", ".log", ".json"):
            return self.load_txt(file_bytes, source=filename)
        else:
            raise ValueError(
                f"Unsupported file type '{suffix}'. "
                "Supported: .pdf, .docx, .txt, .md, .csv, .log, .json"
            )
